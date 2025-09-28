
"""
Fixed data loader with better error handling and fallback for text files.
"""

import logging
from pathlib import Path
from typing import List, Dict, Any
import re

# Try to import python-docx, fallback to text processing if not available
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

def preprocess_text(text: str) -> str:
    """
    Clean and preprocess text by removing extra spaces and newlines.
    
    Args:
        text: Input text to preprocess
        
    Returns:
        Cleaned text string
    """
    if not text:
        return ""
    
    # Remove extra whitespace and normalize
    text = re.sub(r'\s+', ' ', text.strip())
    return text

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
    """
    Split text into chunks while respecting sentence boundaries when possible.
    
    Args:
        text: Input text to chunk
        chunk_size: Maximum size of each chunk
        overlap: Number of characters to overlap between chunks
        
    Returns:
        List of text chunks
    """
    if not text:
        return []
    
    chunks = []
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    current_chunk = ""
    current_length = 0
    
    for sentence in sentences:
        sentence_length = len(sentence)
        
        if current_length + sentence_length <= chunk_size:
            current_chunk += sentence + " "
            current_length += sentence_length + 1
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            
            # Handle overlap
            overlap_text = current_chunk[-overlap:] if overlap > 0 and len(current_chunk) > overlap else ""
            current_chunk = overlap_text + sentence + " "
            current_length = len(overlap_text) + sentence_length + 1
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks

def extract_table_content(table: Any, doc_name: str, table_index: int) -> List[Dict[str, Any]]:
    """
    Extract content from a Word document table.
    
    Args:
        table: Document table object
        doc_name: Name of the document
        table_index: Index of the table in the document
        
    Returns:
        List of table row items with metadata
    """
    items = []
    if not hasattr(table, 'rows') or not table.rows:
        return items
    
    try:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        for row_idx, row in enumerate(table.rows[1:], 1):
            row_text = ""
            for idx, cell in enumerate(row.cells):
                if idx < len(headers):
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text += f"{headers[idx]}: {cell_text} "
            
            if row_text.strip():
                items.append({
                    'id': f"table_{table_index}_row_{row_idx}",
                    'type': 'table_row',
                    'text': preprocess_text(row_text),
                    'metadata': {
                        'source': doc_name,
                        'table_index': table_index,
                        'row_index': row_idx
                    }
                })
    except Exception as e:
        logger.warning(f"Error extracting table content: {e}")
    
    return items

def load_docx(doc_path: Path) -> List[Dict[str, Any]]:
    """
    Load content from a single Word document with fallback to text processing.
    
    Args:
        doc_path: Path to the .docx file
        
    Returns:
        List of items (paragraphs and table rows) with metadata
    """
    if not doc_path.exists():
        logger.error(f"Document not found: {doc_path}")
        raise FileNotFoundError(f"Document not found: {doc_path}")
    
    items = []
    
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not available, trying text file fallback")
        return load_text_file(doc_path)
    
    try:
        doc = Document(str(doc_path))
        
        # Process paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs, 1):
            text = paragraph.text.strip()
            if text and len(text) > 10:  # Skip very short paragraphs
                items.append({
                    'id': f"para_{para_idx}",
                    'type': 'paragraph',
                    'text': preprocess_text(text),
                    'metadata': {
                        'source': doc_path.name,
                        'paragraph_index': para_idx
                    }
                })
        
        # Process tables
        for table_idx, table in enumerate(doc.tables, 1):
            table_items = extract_table_content(table, doc_path.name, table_idx)
            items.extend(table_items)
        
        logger.info(f"Loaded {len(items)} items from {doc_path.name}")
        return items
    
    except Exception as e:
        logger.error(f"Error loading document {doc_path.name}: {e}")
        # Try text file fallback
        logger.info("Attempting text file fallback...")
        return load_text_file(doc_path)

def load_text_file(file_path: Path) -> List[Dict[str, Any]]:
    """
    Load content from a text file as fallback.
    
    Args:
        file_path: Path to the text file
        
    Returns:
        List of items with metadata
    """
    items = []
    
    try:
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        else:
            logger.error(f"Could not decode file {file_path}")
            return items
        
        # Split content into sections or paragraphs
        sections = re.split(r'\n\s*\n', content)
        
        for idx, section in enumerate(sections, 1):
            section = section.strip()
            if section and len(section) > 20:  # Skip very short sections
                items.append({
                    'id': f"section_{idx}",
                    'type': 'text_section',
                    'text': preprocess_text(section),
                    'metadata': {
                        'source': file_path.name,
                        'section_index': idx,
                        'file_type': 'text'
                    }
                })
        
        logger.info(f"Loaded {len(items)} text sections from {file_path.name}")
        return items
        
    except Exception as e:
        logger.error(f"Error loading text file {file_path}: {e}")
        return []

def load_multiple_docx(data_dir: Path) -> List[Dict[str, Any]]:
    """
    Load multiple documents from a directory with fallback support.
    
    Args:
        data_dir: Directory containing document files
        
    Returns:
        Combined list of items from all documents
    """
    all_items = []
    
    # Look for .docx files first
    docx_files = list(data_dir.glob("*.docx")) if data_dir.exists() else []
    
    # If no .docx files, look for text files
    if not docx_files and data_dir.exists():
        text_files = list(data_dir.glob("*.txt"))
        logger.info(f"No .docx files found, using {len(text_files)} text files")
        docx_files = text_files
    
    if not docx_files:
        logger.warning(f"No document files found in {data_dir}")
        return all_items
    
    logger.info(f"Found {len(docx_files)} files to process")
    
    for i, doc_file in enumerate(docx_files, 1):
        try:
            logger.info(f"Processing file {i}/{len(docx_files)}: {doc_file.name}")
            
            if doc_file.suffix.lower() == '.docx':
                items = load_docx(doc_file)
            else:
                items = load_text_file(doc_file)
            
            # Add document source info to metadata
            for item in items:
                item['metadata']['document_name'] = doc_file.name
                item['metadata']['document_index'] = i
                item['id'] = f"{doc_file.stem}_{item['id']}"  # Prefix with filename
            
            all_items.extend(items)
            logger.info(f"Extracted {len(items)} items from {doc_file.name}")
            
        except Exception as e:
            logger.error(f"Error processing {doc_file.name}: {e}")
            continue
    
    logger.info(f"Total items extracted from all documents: {len(all_items)}")
    return all_items

def get_document_summary(data_dir: Path) -> Dict[str, Any]:
    """
    Get summary of all documents in the directory.
    
    Args:
        data_dir: Directory containing documents
        
    Returns:
        Summary statistics
    """
    if not data_dir.exists():
        return {
            'total_documents': 0,
            'document_names': [],
            'total_size_mb': 0,
            'documents': []
        }
    
    # Look for both .docx and .txt files
    all_files = list(data_dir.glob("*.docx")) + list(data_dir.glob("*.txt"))
    
    summary = {
        'total_documents': len(all_files),
        'document_names': [f.name for f in all_files],
        'total_size_mb': sum(f.stat().st_size for f in all_files) / (1024 * 1024),
        'documents': []
    }
    
    for doc_file in all_files:
        try:
            file_info = {
                'name': doc_file.name,
                'size_mb': doc_file.stat().st_size / (1024 * 1024),
                'type': doc_file.suffix,
                'paragraphs': 0,
                'tables': 0,
                'estimated_items': 0
            }
            
            if doc_file.suffix.lower() == '.docx' and DOCX_AVAILABLE:
                # Quick analysis for .docx files
                try:
                    doc = Document(str(doc_file))
                    para_count = len([p for p in doc.paragraphs if p.text.strip()])
                    table_count = len(doc.tables)
                    
                    file_info.update({
                        'paragraphs': para_count,
                        'tables': table_count,
                        'estimated_items': para_count + sum(len(table.rows) for table in doc.tables)
                    })
                except Exception as e:
                    logger.warning(f"Could not analyze .docx file {doc_file.name}: {e}")
                    file_info['estimated_items'] = 'unknown'
            
            else:
                # Quick analysis for text files
                try:
                    with open(doc_file, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    sections = len(re.split(r'\n\s*\n', content))
                    file_info.update({
                        'paragraphs': sections,
                        'tables': 0,
                        'estimated_items': sections
                    })
                except Exception as e:
                    logger.warning(f"Could not analyze text file {doc_file.name}: {e}")
                    file_info['estimated_items'] = 'unknown'
            
            summary['documents'].append(file_info)
            
        except Exception as e:
            logger.warning(f"Could not analyze {doc_file.name}: {e}")
    
    return summary

def extract_qa_pairs(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Extract question-answer pairs from a list of items.
    
    Args:
        items: List of items with text and metadata
        
    Returns:
        List of Q&A pairs
    """
    qa_pairs = []
    i = 0
    
    while i < len(items) - 1:
        current_text = items[i]['text']
        next_text = items[i + 1]['text']
        
        # Look for Q: and A: patterns
        if (current_text.strip().lower().startswith('q:') and 
            next_text.strip().lower().startswith('a:')):
            
            question = preprocess_text(current_text[2:].strip())
            answer = preprocess_text(next_text[2:].strip())
            
            if question and answer:
                qa_pairs.append({
                    'question': question,
                    'answer': answer,
                    'metadata': {
                        'source': items[i]['metadata'].get('source', 'unknown'),
                        'question_id': items[i]['id'],
                        'answer_id': items[i + 1]['id']
                    }
                })
            
            i += 2  # Skip both Q and A items
        else:
            i += 1
    
    return qa_pairs

def create_fallback_data() -> List[Dict[str, Any]]:
    """
    Create fallback sample data when no documents are available.
    
    Returns:
        List of sample document items
    """
    logger.info("Creating fallback legal document data...")
    
    sample_data = [
        {
            'id': 'fallback_1',
            'type': 'paragraph',
            'text': 'Arbitration proceedings typically take 6-18 months from initiation to final award. Emergency arbitration procedures can be completed within 14-30 days for urgent matters requiring immediate relief.',
            'metadata': {'source': 'fallback_legal_data.txt', 'section': 'timeline'}
        },
        {
            'id': 'fallback_2',
            'type': 'paragraph', 
            'text': 'Arbitration costs typically range from $50,000 to $500,000 depending on claim value and complexity. Cost breakdown includes administrative fees (5-15%), arbitrator fees (60-70%), legal representation (20-30%), and other expenses (5-10%).',
            'metadata': {'source': 'fallback_legal_data.txt', 'section': 'costs'}
        },
        {
            'id': 'fallback_3',
            'type': 'paragraph',
            'text': 'For arbitration clauses to be valid and enforceable: clear and unambiguous language required, mutual agreement by all parties, proper scope definition, designated arbitration rules and institution, seat/venue specified, governing law identified.',
            'metadata': {'source': 'fallback_legal_data.txt', 'section': 'validity'}
        },
        {
            'id': 'fallback_4',
            'type': 'paragraph',
            'text': 'Emergency arbitration is available when immediate harm would occur without relief, regular arbitration process would be too slow, interim measures are necessary to preserve status quo, and irreparable damage is imminent.',
            'metadata': {'source': 'fallback_legal_data.txt', 'section': 'emergency'}
        },
        {
            'id': 'fallback_5',
            'type': 'paragraph',
            'text': 'Arbitral awards are generally final and binding, with limited grounds for challenge: procedural irregularities, arbitrator misconduct or bias, award exceeds scope of arbitration agreement, award violates public policy.',
            'metadata': {'source': 'fallback_legal_data.txt', 'section': 'enforcement'}
        }
    ]
    
    return sample_data
