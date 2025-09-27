"""
Creates sample case study data for testing and demonstration.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches
import logging

def create_sample_case_studies():
    """Create a sample Word document with case studies."""
    
    # Create data directory
    data_dir = Path("data")
    data_dir.mkdir(exist_ok=True)
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Business Case Studies Collection', 0)
    
    # Case Study 1: Customer Service
    doc.add_heading('Case Study 1: TechCorp Customer Service Transformation', level=1)
    
    doc.add_paragraph(
        "Problem: TechCorp, a software company with 500 employees, faced declining customer satisfaction "
        "scores dropping from 85% to 62% over 18 months. The main issues included slow response times "
        "averaging 48 hours for email support, lack of 24/7 support coverage, and insufficient training "
        "for support staff on new products."
    )
    
    doc.add_paragraph(
        "Solution: The company implemented a comprehensive customer service overhaul including: "
        "1) Automated ticket routing system reducing response time to 4 hours, "
        "2) Hired additional support staff and implemented 24/7 coverage, "
        "3) Developed intensive 40-hour training program for all support representatives, "
        "4) Introduced customer feedback loops and regular satisfaction surveys."
    )
    
    doc.add_paragraph(
        "Result: Within 6 months, customer satisfaction scores improved to 89%, exceeding the original "
        "baseline. Response times decreased by 92%, and customer retention increased by 23%. "
        "The total investment of $180,000 generated an estimated ROI of 340% through reduced churn."
    )
    
    # Case Study 2: Employee Retention
    doc.add_heading('Case Study 2: GlobalManufacturing Employee Retention Initiative', level=1)
    
    doc.add_paragraph(
        "Problem: GlobalManufacturing experienced high employee turnover at 45% annually, "
        "significantly above the industry average of 18%. Exit interviews revealed key issues: "
        "limited career advancement opportunities, inadequate compensation compared to competitors, "
        "poor work-life balance, and lack of professional development programs."
    )
    
    doc.add_paragraph(
        "Solution: A comprehensive retention strategy was implemented: "
        "1) Created clear career progression paths with defined milestones, "
        "2) Conducted market analysis and adjusted salaries to competitive levels, "
        "3) Introduced flexible working arrangements including remote work options, "
        "4) Established $50,000 annual budget for employee training and certifications, "
        "5) Implemented mentorship program pairing senior and junior employees."
    )
    
    doc.add_paragraph(
        "Result: Employee turnover decreased to 15% within 12 months, saving an estimated "
        "$2.1 million in recruitment and training costs. Employee satisfaction scores increased "
        "from 3.2 to 4.6 out of 5. Productivity improved by 28% due to reduced disruption from turnover."
    )
    
    # Case Study 3: Digital Transformation
    doc.add_heading('Case Study 3: RetailPlus Digital Transformation', level=1)
    
    doc.add_paragraph(
        "Problem: RetailPlus, a traditional retail chain with 200 stores, faced declining sales "
        "due to e-commerce competition. Online sales represented only 5% of total revenue, "
        "while customer foot traffic decreased by 35% over two years. Legacy systems hindered "
        "inventory management and customer experience."
    )
    
    doc.add_paragraph(
        "Solution: Comprehensive digital transformation initiative: "
        "1) Developed omnichannel e-commerce platform with mobile app, "
        "2) Implemented integrated inventory management system across all channels, "
        "3) Introduced click-and-collect services and curbside pickup, "
        "4) Created personalized marketing campaigns using customer data analytics, "
        "5) Trained 800+ staff on digital tools and customer experience standards."
    )
    
    doc.add_paragraph(
        "Result: Online sales grew to 35% of total revenue within 18 months. "
        "Overall sales increased by 22% despite reduced physical footfall. "
        "Customer satisfaction improved to 4.3/5 stars. The $2.8 million investment "
        "achieved payback within 14 months."
    )
    
    # Case Study 4: Cost Reduction
    doc.add_heading('Case Study 4: ManufacturingCorp Operational Efficiency', level=1)
    
    doc.add_paragraph(
        "Problem: ManufacturingCorp faced rising operational costs threatening profitability. "
        "Energy costs increased 40% year-over-year, waste disposal costs rose 25%, "
        "and equipment maintenance consumed 18% of revenue. Production efficiency "
        "declined due to frequent equipment downtime and supply chain disruptions."
    )
    
    doc.add_paragraph(
        "Solution: Systematic cost reduction and efficiency improvement program: "
        "1) Implemented predictive maintenance using IoT sensors reducing downtime by 60%, "
        "2) Introduced lean manufacturing principles eliminating 30% of waste, "
        "3) Negotiated new energy contracts and installed solar panels reducing energy costs by 35%, "
        "4) Diversified supplier base and implemented just-in-time inventory management, "
        "5) Automated repetitive processes reducing labor costs by $400,000 annually."
    )
    
    doc.add_paragraph(
        "Result: Total operational costs reduced by 28% within one year. "
        "Production efficiency improved by 45% through reduced downtime and optimized processes. "
        "Quality defects decreased by 52% due to better maintenance and process controls. "
        "Annual savings of $1.8 million with initial investment of $650,000."
    )
    
    # Case Study 5: Training Program
    doc.add_heading('Case Study 5: FinanceFirst Leadership Development Program', level=1)
    
    doc.add_paragraph(
        "Problem: FinanceFirst identified a leadership pipeline gap with 65% of managers "
        "approaching retirement within 5 years. Internal promotion rates were low at 35%, "
        "and external leadership hires had a 40% failure rate within two years. "
        "Employee engagement scores among mid-level staff were declining."
    )
    
    doc.add_paragraph(
        "Solution: Comprehensive leadership development initiative: "
        "1) Created 18-month leadership development program for high-potential employees, "
        "2) Established executive coaching and 360-degree feedback systems, "
        "3) Implemented cross-functional project assignments and job rotation, "
        "4) Developed internal mentoring network with senior executives, "
        "5) Introduced leadership competency assessments and individualized development plans."
    )
    
    doc.add_paragraph(
        "Result: Internal promotion rate increased to 78% within two years. "
        "Leadership pipeline strength improved with 85% of participants ready for promotion. "
        "Employee engagement scores increased by 32% among program participants. "
        "External hiring costs reduced by $800,000 annually due to successful internal promotions."
    )
    
    # Add Q&A Section
    doc.add_heading('Common Questions and Answers', level=1)
    
    doc.add_paragraph("Q: What are the most common problems companies face?")
    doc.add_paragraph(
        "A: Based on these case studies, the most frequent issues include customer service "
        "challenges, employee retention problems, operational inefficiencies, digital transformation "
        "needs, and leadership development gaps."
    )
    
    doc.add_paragraph("Q: How long do improvement initiatives typically take?")
    doc.add_paragraph(
        "A: Most successful initiatives show measurable results within 6-18 months, "
        "with full transformation typically achieved within 2-3 years depending on scope and complexity."
    )
    
    doc.add_paragraph("Q: What factors contribute to successful implementation?")
    doc.add_paragraph(
        "A: Key success factors include strong leadership commitment, adequate resource allocation, "
        "clear communication strategies, employee training and support, and regular monitoring "
        "with performance metrics."
    )
    
    doc.add_paragraph("Q: How do companies measure ROI on improvement initiatives?")
    doc.add_paragraph(
        "A: ROI is typically measured through cost savings, revenue increases, efficiency gains, "
        "reduced turnover costs, and improved customer satisfaction leading to higher retention rates."
    )
    
    # Save document
    doc_path = data_dir / "cases.docx"
    doc.save(str(doc_path))
    
    print(f"✅ Sample case studies created at: {doc_path}")
    return doc_path

def create_batch_data():
    """Create additional batch data for testing."""
    batch_dir = Path("data") / "batch_1"
    batch_dir.mkdir(exist_ok=True)
    
    # Create additional case study document
    doc = Document()
    
    doc.add_heading('Advanced Business Case Studies', 0)
    
    # Case Study 6: Quality Management
    doc.add_heading('Case Study 6: QualityTech Six Sigma Implementation', level=1)
    
    doc.add_paragraph(
        "Problem: QualityTech manufacturing faced quality control issues with 8.5% defect rates, "
        "customer complaints increased 150% year-over-year, and warranty claims cost $2.3 million annually. "
        "Production inconsistencies led to delayed deliveries and customer dissatisfaction."
    )
    
    doc.add_paragraph(
        "Solution: Implemented comprehensive Six Sigma program: "
        "1) Trained 50+ employees in Six Sigma methodologies, "
        "2) Established quality control checkpoints at each production stage, "
        "3) Implemented statistical process control using real-time monitoring, "
        "4) Created cross-functional quality improvement teams, "
        "5) Introduced supplier quality audits and certification requirements."
    )
    
    doc.add_paragraph(
        "Result: Defect rates reduced to 1.2% within 8 months. "
        "Customer complaints decreased by 85%, warranty claims reduced by 75%. "
        "Production efficiency improved 31% through reduced rework. "
        "Annual savings of $1.9 million achieved."
    )
    
    # Case Study 7: Innovation Management
    doc.add_heading('Case Study 7: InnovaCorp Innovation Lab Initiative', level=1)
    
    doc.add_paragraph(
        "Problem: InnovaCorp struggled with declining market share as competitors "
        "introduced innovative products faster. R&D spending was high but produced "
        "few commercially viable products. Employee suggestions for improvements "
        "were not systematically captured or evaluated."
    )
    
    doc.add_paragraph(
        "Solution: Established comprehensive innovation ecosystem: "
        "1) Created dedicated innovation lab with cross-functional teams, "
        "2) Implemented idea management platform for employee suggestions, "
        "3) Introduced innovation time policy allowing 10% time for creative projects, "
        "4) Established partnerships with universities and startups, "
        "5) Created innovation metrics and reward systems."
    )
    
    doc.add_paragraph(
        "Result: Launched 12 new products within 18 months, generating $15 million revenue. "
        "Employee innovation participation increased to 67%. "
        "Time-to-market for new products reduced by 40%. "
        "Innovation lab achieved 3:1 ROI in first year."
    )
    
    # Case Study 8: Supply Chain Optimization
    doc.add_heading('Case Study 8: LogisticsPro Supply Chain Transformation', level=1)
    
    doc.add_paragraph(
        "Problem: LogisticsPro faced supply chain disruptions causing 25% delivery delays, "
        "inventory holding costs increased 40%, and customer complaints about late deliveries "
        "rose 200%. Limited visibility into supplier performance and inventory levels "
        "hindered decision-making and responsiveness."
    )
    
    doc.add_paragraph(
        "Solution: Comprehensive supply chain digitization and optimization: "
        "1) Implemented end-to-end supply chain visibility platform, "
        "2) Introduced predictive analytics for demand forecasting, "
        "3) Established strategic partnerships with key suppliers, "
        "4) Created automated inventory management with dynamic safety stock levels, "
        "5) Deployed IoT sensors for real-time shipment tracking."
    )
    
    doc.add_paragraph(
        "Result: On-time delivery improved to 96% within 12 months. "
        "Inventory holding costs reduced by 32% through optimized stock levels. "
        "Supply chain visibility increased to 99% across all touchpoints. "
        "Customer satisfaction scores for delivery improved from 2.8 to 4.7 out of 5."
    )
    
    # Save batch document
    batch_doc_path = batch_dir / "advanced_cases.docx"
    doc.save(str(batch_doc_path))
    
    print(f"✅ Batch case studies created at: {batch_doc_path}")
    return batch_doc_path

def create_simple_fallback_data():
    """Create simple text-based fallback data if docx creation fails."""
    data_dir = Path("data")
    data_dir.mkdir(exist_ok=True)
    
    # Create simple text file with case studies
    fallback_content = """BUSINESS CASE STUDIES COLLECTION

=== Case Study 1: Customer Service Improvement ===
Problem: Company faced 48-hour email response times and declining satisfaction scores.
Solution: Automated routing, 24/7 coverage, staff training, and feedback systems.
Result: Satisfaction improved from 62% to 89%, response time reduced to 4 hours, ROI of 340%.

=== Case Study 2: Employee Retention ===
Problem: High 45% annual turnover vs 18% industry average.
Solution: Career paths, competitive salaries, flexible work, training budget, mentorship.
Result: Turnover reduced to 15%, saved $2.1M in recruitment costs.

=== Case Study 3: Digital Transformation ===
Problem: Declining retail sales, only 5% online revenue, legacy systems.
Solution: E-commerce platform, inventory integration, click-and-collect, staff training.
Result: Online sales grew to 35%, overall sales up 22%, ROI in 14 months.

=== Case Study 4: Cost Reduction ===
Problem: Rising operational costs, energy up 40%, maintenance consuming 18% of revenue.
Solution: Predictive maintenance, lean manufacturing, energy optimization, supplier diversification.
Result: Total costs reduced 28%, efficiency up 45%, $1.8M annual savings.

=== Case Study 5: Leadership Development ===
Problem: 65% of managers retiring soon, low internal promotion rates.
Solution: 18-month development program, coaching, job rotation, mentoring network.
Result: Internal promotions up to 78%, engagement increased 32%.

=== Common Q&A ===
Q: What are the most common business problems?
A: Customer service issues, employee retention, operational inefficiencies, digital transformation needs.

Q: How long do improvements take?
A: Most show results in 6-18 months, full transformation in 2-3 years.

Q: What ensures success?
A: Leadership commitment, adequate resources, clear communication, training, regular monitoring.
"""
    
    fallback_path = data_dir / "fallback_cases.txt"
    fallback_path.write_text(fallback_content)
    
    print(f"✅ Fallback case studies created at: {fallback_path}")
    return fallback_path

def main():
    """Create all sample data."""
    print("🚀 Creating sample case study data...")
    
    try:
        # Try to create Word documents first
        main_doc = create_sample_case_studies()
        batch_doc = create_batch_data()
        
        print(f"\n✅ Sample data creation complete!")
        print(f"📁 Main document: {main_doc}")
        print(f"📁 Batch document: {batch_doc}")
        print(f"\nYou can now run the chatbot with sample data.")
        
    except ImportError as e:
        print(f"⚠️ python-docx not available: {e}")
        print("📝 Creating simple text-based fallback data...")
        fallback_doc = create_simple_fallback_data()
        print(f"✅ Fallback data created: {fallback_doc}")
        print(f"\nInstall python-docx for full Word document support:")
        print(f"pip install python-docx")
        
    except Exception as e:
        print(f"❌ Error creating sample data: {e}")
        print("📝 Creating simple fallback data...")
        try:
            fallback_doc = create_simple_fallback_data()
            print(f"✅ Fallback data created: {fallback_doc}")
        except Exception as fallback_error:
            print(f"❌ Fallback creation also failed: {fallback_error}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    main()